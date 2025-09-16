from docx import Document
from docx.shared import Inches

def create_mom_docx(mom_data, filename="meeting_minutes.docx"):
    """Create a DOCX file from MoM JSON data"""
    doc = Document()
    
    # Title
    title = doc.add_heading(mom_data.get('meeting_title', 'Meeting Minutes'), 0)
    
    # Date
    doc.add_paragraph(f"Date: {mom_data.get('date', 'Unknown')}")
    
    # Attendees
    doc.add_heading('Attendees', level=1)
    for attendee in mom_data.get('attendees', []):
        doc.add_paragraph(f"• {attendee}")
    
    # Agenda Items
    if mom_data.get('agenda_items'):
        doc.add_heading('Agenda Items', level=1)
        for item in mom_data['agenda_items']:
            doc.add_paragraph(f"• {item}")
    
    # Key Points
    if mom_data.get('key_points'):
        doc.add_heading('Key Points', level=1)
        for point in mom_data['key_points']:
            doc.add_paragraph(f"• {point}")
    
    # Action Items
    if mom_data.get('action_items'):
        doc.add_heading('Action Items', level=1)
        for action in mom_data['action_items']:
            doc.add_paragraph(f"• {action['task']} (Owner: {action['owner']}, Due: {action['due_date']})")
    
    # Decisions Made
    if mom_data.get('decisions_made'):
        doc.add_heading('Decisions Made', level=1)
        for decision in mom_data['decisions_made']:
            doc.add_paragraph(f"• {decision}")
    
    # Next Steps
    if mom_data.get('next_steps'):
        doc.add_heading('Next Steps', level=1)
        for step in mom_data['next_steps']:
            doc.add_paragraph(f"• {step}")
    
    doc.save(filename)
    return filename